from django.shortcuts import render,get_object_or_404,redirect,HttpResponse
from django.contrib.auth.models import User
from django.views.generic import ListView, DetailView,CreateView,UpdateView, DeleteView
from .models import Post,Category,Comment
from django.contrib.auth.mixins import LoginRequiredMixin
from .forms import PostForms, EditForms, CommentForm
from django.urls import reverse
from django.views.generic import View
from django.http import HttpResponseRedirect
from django.http import HttpResponseForbidden
from django.http import JsonResponse
from django.core.exceptions import PermissionDenied
from django.db.models import Q 
from docx import Document
from PIL import Image
from docx.shared import Inches
from io import BytesIO
from django.db.models import Count



from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from django.urls import reverse_lazy
# Create your views here.
 #def home(request):
   # return render(request, 'Home.html',{})

from django.contrib.auth.decorators import login_required
from .models import Profile

def follow_user(request, pk):
    if request.method == 'POST' and request.user.is_authenticated:
        profile_to_follow = get_object_or_404(Profile, pk=pk)
        if request.user != profile_to_follow.user:
            profile_to_follow.followers.add(request.user)
            return JsonResponse({'status': 'followed', 'followers_count': profile_to_follow.followers.count()})
    return JsonResponse({'status': 'error'}, status=400)

def unfollow_user(request, pk):
    if request.method == 'POST' and request.user.is_authenticated:
        profile_to_unfollow = get_object_or_404(Profile, pk=pk)
        if request.user != profile_to_unfollow.user:
            profile_to_unfollow.followers.remove(request.user)
            return JsonResponse({'status': 'unfollowed', 'followers_count': profile_to_unfollow.followers.count()})
    return JsonResponse({'status': 'error'}, status=400)




def like_posts(request,pk):
    post = get_object_or_404(Post,id=pk)
    if request.method == 'POST':
        if request.user in post.likes.all():  # If the user already liked the post, unlike it
            post.likes.remove(request.user)
        else:  # Otherwise, like the post
            post.likes.add(request.user)
        return HttpResponseRedirect(reverse('article-detail', args=[str(pk)]))


class HomeView(ListView):
    model = Post
    template_name = 'home.html'
    ordering = ['-add_date']
    paginate_by = 6

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['cat_menu'] = Category.objects.all()

        # Annotate posts with like counts and order by like count in descending order
        trending_posts = Post.objects.annotate(like_count=Count('likes')).order_by('-like_count')[:3]
        context['trending_posts'] = trending_posts

        return context

    
       
class ArticleDetailView(DetailView):
    model = Post
    template_name='Article_detail.html' 
    context_object_name = 'post'
    def get_context_data(self,*args, **kwargs):
        cat_menu= Category.objects.all()
        context= super(ArticleDetailView,self).get_context_data(*args, **kwargs)
        stuff= get_object_or_404(Post, id=self.kwargs['pk'])
        total_likes= stuff.total_likes()
        context["cat_menu"] = cat_menu
        context["total_likes"] = total_likes
        context['comments'] = Comment.objects.filter(post=self.object).order_by('-date_added')
        return context


class AddPostView(CreateView):
    model = Post
    form_class = PostForms
    template_name = 'add_post.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = Category.objects.all()  # Pass all categories to the template
        return context

    def form_valid(self, form):
        title = form.cleaned_data['title']
        title_tag = form.cleaned_data['title_tag']
        
        if not title.isalnum():
            form.add_error('title', 'Title should be alphanumeric.')
            return self.form_invalid(form)
        
        if not title_tag.isalpha():
            form.add_error('title_tag', 'Title-Tag should only consist of alphabets.')
            return self.form_invalid(form)
        post = form.save(commit=False)  # Get the unsaved instance of the post form
        post.author = self.request.user  # Assign the author to the current user
        post.save()  # Save the post instance to the database
        
        return super().form_valid(form)



class AddCommentView(CreateView):
    model = Comment
    form_class= CommentForm
    template_name='add_comment.html'
    #fields = '__all__'  
    def form_valid(self, form):
        form.instance.post_id = self.kwargs['pk']
        form.instance.user = self.request.user
        return super().form_valid(form)
    
    def get_success_url(self):
        # Redirect to the detail view of the post where the comment was added
        return reverse_lazy('article-detail', kwargs={'pk': self.kwargs['pk']})


class DeleteCommentView(LoginRequiredMixin, DeleteView):
    model = Comment
    template_name = 'delete_comment.html'

    def get_queryset(self):
        queryset = super().get_queryset()
        return queryset.filter(user=self.request.user) | queryset.filter(post__author=self.request.user)

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        if self.object.user == request.user or self.object.post.author == request.user:
            return super().delete(request, *args, **kwargs)
        else:
            return HttpResponseForbidden("You are not allowed to delete this comment.")

    def get_success_url(self):
        return reverse_lazy('article-detail', kwargs={'pk': self.object.post.pk})
    
    
        

class LikeCommentView(View):
    def post(self, request, pk):
        # Get the comment object or return a 404 error if not found
        comment = get_object_or_404(Comment, pk=pk)
        
        # Check if the user has already liked the comment
        if request.user.is_authenticated:
            if request.user in comment.likes.all():
                # If the user has already liked the comment, unlike it
                comment.likes.remove(request.user)
                liked = False
            else:
                # Otherwise, like the comment
                comment.likes.add(request.user)
                liked = True

            # Return a JSON response with the updated like count and liked status
            return JsonResponse({'liked': liked, 'like_count': comment.likes.count()})
        else:
            # Return a JSON response indicating that the user must be logged in to like comments
            return JsonResponse({'error': 'User must be logged in to like comments'}, status=403)

class AddCategoryView(CreateView):
    model = Category
    #form_class= PostForms
    template_name='add_category.html'
    fields = '__all__'    
    success_url = reverse_lazy('Home')


class UpdatePostView(LoginRequiredMixin, UpdateView):
    model = Post
    form_class = EditForms
    template_name = 'update_post.html'

    def get_queryset(self):
        queryset = super().get_queryset()
        return queryset.filter(author=self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['form'] = self.get_form()  # Add this line to include the form in the context
        return context

    def get_object(self, queryset=None):
        obj = super().get_object(queryset)
        if obj.author != self.request.user:
            raise PermissionDenied("You do not have permission to edit this post.")
        return obj

    def get_success_url(self):
        return reverse_lazy('article-detail', kwargs={'pk': self.object.pk})
        

class DeletePostView(LoginRequiredMixin, DeleteView):
    model = Post
    template_name = 'delete_post.html'
    success_url = reverse_lazy('Home')

    def get_queryset(self):
        queryset = super().get_queryset()
        # Ensure the user can only delete their own posts
        return queryset.filter(author=self.request.user)

    def get(self, request, *args, **kwargs):
        # Check if the user is authenticated
        if not request.user.is_authenticated:
            return self.handle_no_permission()
        # If the user is authenticated, call the parent get method
        return super().get(request, *args, **kwargs)

class CategoryView(ListView):
    model = Post
    template_name = 'categories.html'
    context_object_name = 'cat_posts'
    paginate_by = 6

    def get_queryset(self):
        category_title = self.kwargs['category_title']
        return Post.objects.filter(cat__title=category_title)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['cat'] = self.kwargs['category_title']
        return context
    

def CategoryListView(request):
    cat_menu_list= Category.objects.all()
    return render(request, 'category_list.html', {'cat_menu_list': cat_menu_list })

def Search(request):
    query = request.GET.get('query')
    if query:
        # Check if the query matches a category title
        categories = Category.objects.filter(title__icontains=query)
        if categories:
            posts = Post.objects.filter(cat__in=categories)
            authors = User.objects.none()  # No authors for category search
        else:
            posts = Post.objects.filter(
                Q(title__icontains=query) | 
                Q(content__icontains=query) | 
                Q(author__first_name__icontains=query) | 
                Q(author__last_name__icontains=query) |
                Q(author__username__icontains=query)
            )
            # Check if the query matches an author's name
            authors = User.objects.filter(
                Q(first_name__icontains=query) | 
                Q(last_name__icontains=query) |
                Q(username__icontains=query) 

            )
        # If no posts found, reset categories and authors
        if not posts:
            categories = Category.objects.none()
            authors = User.objects.none()
    else:
        posts = Post.objects.none()
        categories = Category.objects.none()
        authors = User.objects.none()

    context = {
        'query': query,
        'posts': posts,
        'categories': categories,
        'authors': authors,
    }
    

    return render(request, 'search.html', context)

from docx import Document
from django.http import HttpResponse

def download_word_document(request, post_id):
    # Retrieve the post object
    post = get_object_or_404(Post, id=post_id)
    
    # Create a new Document object
    document = Document()
    
    # Add the post title to the document
    document.add_heading(post.title, level=1)
    
    # Add the post author's name to the document
    document.add_paragraph(f'Author: {post.author.first_name} {post.author.last_name}')
    
    # Add the post image to the document
    if post.img:
        # Open the image file
        img = Image.open(post.img.path)
        # Convert the image to a BytesIO object
        img_io = BytesIO()
        img.save(img_io, format='PNG')
        # Add the image to the document
        document.add_picture(img_io, width=Inches(3))
    
    # Add the post content to the document
    document.add_paragraph(post.content)
    
    # Prepare response with Word document content
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{post.title}.docx"'
    
    # Save the document to the response
    document.save(response)
    
    return response



